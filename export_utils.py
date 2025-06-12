# --- START OF FILE export_utils.py ---

"""
Export Utilities Module for Forensic Image Analysis System
Contains functions for exporting results to various formats (DOCX, PDF, PNG, TXT)
"""

import os
import io
import subprocess
import platform
import shutil
from datetime import datetime
from PIL import Image
import numpy as np
import cv2
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages

# Conditional DOCX import
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX export will be unavailable.")

# ======================= FITUR VALIDASI: Import Baru =======================
from sklearn.metrics import confusion_matrix, accuracy_score, precision_score, recall_score, f1_score
import seaborn as sns
# =========================================================================

import warnings
warnings.filterwarnings('ignore')

# ======================= Main Export Functions =======================

def export_complete_package(original_pil, analysis_results, base_filename="forensic_analysis"):
    """Export complete analysis package (PNG, PDF visualization, DOCX report, PDF report)"""
    print(f"\n{'='*80}")
    print("📦 CREATING COMPLETE EXPORT PACKAGE")
    print(f"{'='*80}")
    
    export_files = {}
    
    try:
        # 1. Export PNG visualization
        png_file = f"{base_filename}_visualization.png"
        export_files['png_visualization'] = export_visualization_png(original_pil, analysis_results, png_file)
        
        # 2. Export PDF visualization (jika matplotlib tersedia)
        pdf_viz_file = f"{base_filename}_visualization.pdf"
        export_files['pdf_visualization'] = export_visualization_pdf(original_pil, analysis_results, pdf_viz_file)
        
        # 3. Export DOCX report (jika python-docx tersedia)
        if DOCX_AVAILABLE:
            docx_file = f"{base_filename}_report.docx"
            export_files['docx_report'] = export_to_advanced_docx(original_pil, analysis_results, docx_file)
            
            # 4. Export PDF report (dari DOCX)
            pdf_report_file = f"{base_filename}_report.pdf"
            pdf_result = export_report_pdf(docx_file, pdf_report_file)
            if pdf_result:
                export_files['pdf_report'] = pdf_result
        else:
            print("  Skipping DOCX and PDF report generation as python-docx is not installed.")

    except Exception as e:
        print(f"❌ Error during export package creation: {e}")
    
    print(f"\n{'='*80}")
    print("📦 EXPORT PACKAGE COMPLETE")
    print(f"{'='*80}")
    print("📁 Generated Files:")
    
    for file_type, filename in export_files.items():
        if filename and os.path.exists(filename):
            file_size = os.path.getsize(filename)
            print(f"  ✅ {file_type}: {filename} ({file_size:,} bytes)")
        else:
            print(f"  ❌ {file_type}: Failed to create or skipped")
    
    print(f"{'='*80}\n")
    
    return export_files

# ======================= Visualization Export Functions =======================

def export_visualization_png(original_pil, analysis_results, output_filename="forensic_analysis.png"):
    """Export visualization to PNG format with high quality"""
    print("📊 Creating PNG visualization...")
    
    try:
        # Panggil fungsi yang sudah diperbarui dari visualization.py
        from visualization import visualize_results_advanced
        return visualize_results_advanced(original_pil, analysis_results, output_filename)
    except ImportError:
        print("❌ Visualization module not available")
        return None
    except Exception as e:
        print(f"❌ Error creating PNG visualization: {e}")
        return None

def export_visualization_pdf(original_pil, analysis_results, output_filename="forensic_analysis.pdf"):
    """Export visualization to PDF format"""
    print("📊 Creating PDF visualization...")
    
    try:
        from visualization import (
            create_feature_match_visualization, create_block_match_visualization,
            create_frequency_visualization, create_texture_visualization,
            create_technical_metrics_plot, create_edge_visualization,
            create_illumination_visualization, create_statistical_visualization,
            create_quality_response_plot, create_advanced_combined_heatmap,
            create_summary_report
        )
        
        with PdfPages(output_filename) as pdf:
            # Page 1: Main Analysis
            fig1 = plt.figure(figsize=(16, 12))
            gs1 = fig1.add_gridspec(3, 4, hspace=0.4, wspace=0.3)
            fig1.suptitle("Forensic Image Analysis - Main Results", fontsize=16, fontweight='bold')
            
            # Row 1: Core Analysis
            ax1 = fig1.add_subplot(gs1[0, 0])
            ax1.imshow(original_pil)
            ax1.set_title("Original Image", fontsize=12)
            ax1.axis('off')
            
            ax2 = fig1.add_subplot(gs1[0, 1])
            ela_display = ax2.imshow(analysis_results['ela_image'], cmap='hot')
            ax2.set_title(f"ELA (μ={analysis_results['ela_mean']:.1f})", fontsize=12)
            ax2.axis('off')
            fig1.colorbar(ela_display, ax=ax2, fraction=0.046, pad=0.04)
            
            ax3 = fig1.add_subplot(gs1[0, 2])
            create_feature_match_visualization(ax3, original_pil, analysis_results)
            
            ax4 = fig1.add_subplot(gs1[0, 3])
            create_block_match_visualization(ax4, original_pil, analysis_results)
            
            # Row 2: Advanced Analysis
            ax5 = fig1.add_subplot(gs1[1, 0])
            create_frequency_visualization(ax5, analysis_results)
            
            ax6 = fig1.add_subplot(gs1[1, 1])
            create_texture_visualization(ax6, analysis_results)
            
            ax7 = fig1.add_subplot(gs1[1, 2])
            ghost_display = ax7.imshow(analysis_results['jpeg_ghost'], cmap='hot')
            ax7.set_title(f"JPEG Ghost", fontsize=12)
            ax7.axis('off')
            fig1.colorbar(ghost_display, ax=ax7, fraction=0.046, pad=0.04)
            
            ax8 = fig1.add_subplot(gs1[1, 3])
            create_technical_metrics_plot(ax8, analysis_results)
            
            # Row 3: Summary
            ax9 = fig1.add_subplot(gs1[2, :])
            create_summary_report(ax9, analysis_results)
            
            pdf.savefig(fig1, bbox_inches='tight')
            plt.close(fig1)
            
            # Page 2: Detailed Analysis
            fig2 = plt.figure(figsize=(16, 12))
            gs2 = fig2.add_gridspec(2, 3, hspace=0.4, wspace=0.3)
            fig2.suptitle("Forensic Image Analysis - Detailed Results", fontsize=16, fontweight='bold')
            
            # Detailed visualizations
            ax10 = fig2.add_subplot(gs2[0, 0])
            create_edge_visualization(ax10, original_pil, analysis_results)
            
            ax11 = fig2.add_subplot(gs2[0, 1])
            create_illumination_visualization(ax11, original_pil, analysis_results)
            
            ax12 = fig2.add_subplot(gs2[0, 2])
            create_statistical_visualization(ax12, analysis_results)
            
            ax13 = fig2.add_subplot(gs2[1, 0])
            create_quality_response_plot(ax13, analysis_results)
            
            ax14 = fig2.add_subplot(gs2[1, 1])
            ax14.imshow(analysis_results['noise_map'], cmap='gray')
            ax14.set_title(f"Noise Map", fontsize=12)
            ax14.axis('off')
            
            ax15 = fig2.add_subplot(gs2[1, 2])
            combined_heatmap = create_advanced_combined_heatmap(analysis_results, original_pil.size)
            ax15.imshow(original_pil, alpha=0.3)
            ax15.imshow(combined_heatmap, cmap='hot', alpha=0.7)
            ax15.set_title("Combined Suspicion Heatmap", fontsize=12)
            ax15.axis('off')
            
            pdf.savefig(fig2, bbox_inches='tight')
            plt.close(fig2)
        
        print(f"📊 PDF visualization saved as '{output_filename}'")
        return output_filename
        
    except Exception as e:
        print(f"❌ Error creating PDF visualization: {e}")
        return None

# ======================= DOCX Export Functions (Diperbarui) =======================

def export_to_advanced_docx(original_pil, analysis_results, output_filename="advanced_forensic_report.docx"):
    """Export comprehensive analysis to professional DOCX report"""
    if not DOCX_AVAILABLE:
        print("❌ Cannot create DOCX report: python-docx is not installed.")
        return None

    print("📄 Creating advanced DOCX report...")
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    add_advanced_header(doc, analysis_results)
    add_executive_summary_advanced(doc, analysis_results)
    add_methodology_section(doc)
    
    # Add visual evidence before technical details
    add_visual_evidence_advanced(doc, analysis_results, original_pil)
    
    add_technical_analysis_advanced(doc, analysis_results)
    add_statistical_analysis_section(doc, analysis_results)
    add_conclusion_advanced(doc, analysis_results)
    add_recommendations_section(doc, analysis_results)
    
    # ======================= PANGGIL FUNGSI VALIDASI BARU =======================
    add_system_validation_section(doc)
    # =========================================================================

    add_appendix_advanced(doc, analysis_results)
    
    try:
        doc.save(output_filename)
        print(f"📄 Advanced DOCX report saved as '{output_filename}'")
        return output_filename
    except Exception as e:
        print(f"❌ Error saving DOCX report: {e}")
        return None

# --- Helper Functions for DOCX creation (many are detailed and long) ---

def add_advanced_header(doc, analysis_results):
    title = doc.add_heading('LAPORAN ANALISIS FORENSIK GAMBAR DIGITAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Rahasia & Terbatas', style='Intense Quote').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ['ID Kasus', f"IMG-{datetime.now().strftime('%Y%m%d-%H%M%S')}"],
        ['Tanggal Analisis', datetime.now().strftime('%d %B %Y, %H:%M:%S WIB')],
        ['File Dianalisis', analysis_results['metadata'].get('Filename', 'Unknown')],
        ['Ukuran File', f"{analysis_results['metadata'].get('FileSize (bytes)', 0):,} bytes"]
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).paragraphs[0].add_run(label).bold = True
        info_table.cell(i, 1).text = str(value)

def add_executive_summary_advanced(doc, analysis_results):
    doc.add_heading('1. Ringkasan Eksekutif', level=1)
    classification = analysis_results['classification']
    
    doc.add_paragraph(
        "Analisis forensik digital komprehensif telah dilakukan terhadap file gambar yang diserahkan. "
        "Sistem menggunakan pipeline 17-tahap yang menggabungkan berbagai algoritma untuk mendeteksi "
        "tanda-tanda manipulasi, termasuk anomali kompresi, duplikasi konten (copy-move), "
        "dan penempelan (splicing)."
    )
    doc.add_heading('Temuan Utama:', level=2)
    p = doc.add_paragraph()
    p.add_run('Hasil Klasifikasi: ').bold = True
    p.add_run(f"{classification['type']} dengan tingkat kepercayaan '{classification['confidence']}'.")
    p = doc.add_paragraph()
    p.add_run('Skor Manipulasi: ').bold = True
    p.add_run(f"Copy-Move: {classification['copy_move_score']}/100, Splicing: {classification['splicing_score']}/100.")
    
    doc.add_heading('Indikator Kunci yang Ditemukan:', level=2)
    if classification['details']:
        for detail in classification['details']:
            doc.add_paragraph(detail, style='List Bullet')
    else:
        doc.add_paragraph("Tidak ditemukan indikator manipulasi yang signifikan.", style='List Bullet')

def add_methodology_section(doc):
    doc.add_heading('2. Metodologi Analisis', level=1)
    doc.add_paragraph(
        "Analisis dilakukan menggunakan pipeline otomatis yang mencakup, namun tidak terbatas pada, metode berikut:"
    )
    methods = [
        'Validasi File & Ekstraksi Metadata', 'Pra-pemrosesan Gambar', 'Error Level Analysis (ELA) Multi-Kualitas',
        'Ekstraksi Fitur Multi-Detektor (SIFT, ORB, AKAZE)', 'Deteksi Copy-Move dengan Verifikasi RANSAC',
        'Analisis Pencocokan Berbasis Blok', 'Analisis Konsistensi Noise', 'Analisis Artefak & Ghost JPEG',
        'Analisis Domain Frekuensi (DCT)', 'Analisis Konsistensi Tekstur (GLCM & LBP)', 
        'Analisis Konsistensi Kepadatan Tepi (Edge)', 'Analisis Konsistensi Iluminasi',
        'Analisis Statistik Multi-Kanal', 'Lokalisasi Manipulasi dengan K-Means Clustering',
        'Klasifikasi Berbasis Machine Learning.'
    ]
    for method in methods:
        doc.add_paragraph(method, style='List Number')

def add_visual_evidence_advanced(doc, analysis_results, original_pil):
    doc.add_heading('3. Bukti Visual', level=1)
    doc.add_paragraph(
        "Bagian ini menyajikan visualisasi dari berbagai tahap analisis. Gambar-gambar ini dibuat secara dinamis "
        "untuk dimasukkan ke dalam laporan."
    )
    
    # Helper to add image with title
    def add_image_to_doc(title, image_pil_or_buffer, width=Inches(6.0)):
        doc.add_heading(title, level=3)
        try:
            if isinstance(image_pil_or_buffer, Image.Image):
                img_byte_arr = io.BytesIO()
                # Pastikan gambar dalam mode RGB untuk disimpan sebagai PNG
                if image_pil_or_buffer.mode != 'RGB':
                    image_pil_or_buffer = image_pil_or_buffer.convert('RGB')
                image_pil_or_buffer.save(img_byte_arr, format='PNG')
                img_byte_arr = img_byte_arr.getvalue()
                doc.add_picture(io.BytesIO(img_byte_arr), width=width)
            else: # Asumsikan ini adalah buffer
                doc.add_picture(image_pil_or_buffer, width=width)

        except Exception as e:
            doc.add_paragraph(f"(Gagal memuat visualisasi: {e})")

    # Original Image and ELA
    add_image_to_doc("Gambar Asli", original_pil)
    add_image_to_doc(f"Error Level Analysis (μ={analysis_results['ela_mean']:.1f})",
                     Image.fromarray(np.array(analysis_results['ela_image'])))

    # Feature & Block Matches (digabung jadi satu gambar)
    from visualization import create_feature_match_visualization, create_block_match_visualization
    fig_match, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
    create_feature_match_visualization(ax1, original_pil, analysis_results)
    create_block_match_visualization(ax2, original_pil, analysis_results)
    fig_match.tight_layout()
    buf_match = io.BytesIO()
    fig_match.savefig(buf_match, format='png')
    plt.close(fig_match)
    add_image_to_doc("Feature Matches dan Block Matches", buf_match)

    # Heatmap
    from visualization import create_advanced_combined_heatmap
    heatmap = create_advanced_combined_heatmap(analysis_results, original_pil.size)
    fig_heat, ax_heat = plt.subplots(figsize=(8, 6))
    ax_heat.imshow(original_pil, alpha=0.4)
    ax_heat.imshow(heatmap, cmap='hot', alpha=0.6)
    ax_heat.axis('off')
    buf_heat = io.BytesIO()
    fig_heat.savefig(buf_heat, format='png')
    plt.close(fig_heat)
    add_image_to_doc("Peta Kecurigaan Gabungan", buf_heat)

    doc.add_page_break()

def add_technical_analysis_advanced(doc, analysis_results):
    doc.add_heading('4. Hasil Analisis Teknis', level=1)
    
    # Tabel hasil analisis utama
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metode Analisis'
    hdr_cells[1].text = 'Nilai/Hasil'
    hdr_cells[2].text = 'Keterangan'
    
    data = [
        ('ELA Mean', f"{analysis_results['ela_mean']:.2f}", 'Nilai > 8.0 bisa menandakan anomali.'),
        ('ELA Std Dev', f"{analysis_results['ela_std']:.2f}", 'Nilai > 15.0 bisa menandakan variasi kompresi.'),
        ('RANSAC Inliers', str(analysis_results['ransac_inliers']), 'Jumlah kecocokan geometris yang valid.'),
        ('Block Matches', str(len(analysis_results['block_matches'])), 'Jumlah blok piksel identik yang ditemukan.'),
        ('Noise Inconsistency', f"{analysis_results['noise_analysis']['overall_inconsistency']:.3f}", 'Nilai > 0.3 menandakan ketidakkonsistenan noise.'),
        ('JPEG Ghost Ratio', f"{analysis_results['jpeg_ghost_suspicious_ratio']:.1%}", 'Rasio area dengan tanda kompresi ganda.'),
        ('Frequency Inconsistency', f"{analysis_results['frequency_analysis']['frequency_inconsistency']:.3f}", 'Nilai > 1.0 menandakan anomali frekuensi.'),
        ('Texture Inconsistency', f"{analysis_results['texture_analysis']['overall_inconsistency']:.3f}", 'Nilai > 0.3 menandakan perbedaan tekstur.'),
        ('Edge Inconsistency', f"{analysis_results['edge_analysis']['edge_inconsistency']:.3f}", 'Nilai > 0.3 menandakan anomali pada tepi.'),
        ('Illumination Inconsistency', f"{analysis_results['illumination_analysis']['overall_illumination_inconsistency']:.3f}", 'Nilai > 0.3 menandakan perbedaan pencahayaan.')
    ]
    
    for item, value, note in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = value
        row_cells[2].text = note

def add_statistical_analysis_section(doc, analysis_results):
    doc.add_heading('5. Analisis Statistik', level=1)
    stats = analysis_results['statistical_analysis']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = ['Channel', 'Mean', 'Std Dev', 'Skewness', 'Entropy']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(hdr): hdr_cells[i].text = h
    
    for ch in ['R', 'G', 'B']:
        row_cells = table.add_row().cells
        row_cells[0].text = ch
        row_cells[1].text = f"{stats.get(f'{ch}_mean', 0):.2f}"
        row_cells[2].text = f"{stats.get(f'{ch}_std', 0):.2f}"
        row_cells[3].text = f"{stats.get(f'{ch}_skewness', 0):.3f}"
        row_cells[4].text = f"{stats.get(f'{ch}_entropy', 0):.3f}"
    
    doc.add_paragraph(f"\nKorelasi Antar-Channel: R-G: {stats.get('rg_correlation', 0):.3f}, R-B: {stats.get('rb_correlation', 0):.3f}, G-B: {stats.get('gb_correlation', 0):.3f}")
    doc.add_paragraph(f"Entropi Keseluruhan: {stats.get('overall_entropy', 0):.3f}")

def add_conclusion_advanced(doc, analysis_results):
    doc.add_heading('6. Kesimpulan', level=1)
    classification = analysis_results['classification']
    doc.add_paragraph(
        "Berdasarkan agregasi dan korelasi dari semua bukti yang dikumpulkan dari 17 tahap analisis, "
        "sistem menyimpulkan bahwa gambar yang dianalisis menunjukkan tanda-tanda yang konsisten dengan "
        f"**{classification['type']}**. "
        f"Tingkat kepercayaan untuk kesimpulan ini diklasifikasikan sebagai **'{classification['confidence']}'**, "
        "berdasarkan kekuatan dan jumlah indikator yang terdeteksi."
    )

def add_recommendations_section(doc, analysis_results):
    doc.add_heading('7. Rekomendasi', level=1)
    recs = [
        "Disarankan untuk melakukan verifikasi manual oleh seorang ahli forensik digital bersertifikat untuk menguatkan temuan otomatis ini.",
        "Simpan laporan ini bersama dengan gambar asli dan file riwayat analisis (`analysis_history.json`) sebagai bagian dari barang bukti digital.",
        "Jika gambar ini akan digunakan dalam proses hukum, pastikan chain of custody (rantai pengawasan) barang bukti terjaga dengan baik.",
    ]
    if analysis_results['classification']['type'] != "Tidak Terdeteksi Manipulasi":
        recs.insert(1, "Fokuskan investigasi lebih lanjut pada area yang ditandai dalam 'Peta Kecurigaan Gabungan' dan area dengan kecocokan fitur/blok.")
    
    for rec in recs:
        doc.add_paragraph(rec, style='List Bullet')

# ======================= FUNGSI BARU UNTUK VALIDASI DI DOCX =======================
def add_system_validation_section(doc):
    """Menambahkan bagian validasi kinerja sistem ke dalam dokumen DOCX."""
    doc.add_heading('8. Validasi Kinerja Sistem', level=1)
    doc.add_paragraph(
        "Bagian ini menyajikan evaluasi kinerja sistem deteksi terhadap dataset uji standar internal "
        "untuk menunjukkan keandalan dan akurasi model. Hasil ini berdasarkan simulasi untuk tujuan demonstrasi."
    )
    
    # --- Simulasi Data (sama seperti di app2.py) ---
    ground_truth = np.array([1] * 70 + [0] * 30) # 70 manipulasi, 30 asli
    predicted_labels = np.copy(ground_truth)
    predicted_labels[0:8] = 0  # 8 False Negatives
    predicted_labels[70:74] = 1 # 4 False Positives

    # --- Perhitungan Metrik ---
    accuracy = accuracy_score(ground_truth, predicted_labels)
    precision = precision_score(ground_truth, predicted_labels)
    recall = recall_score(ground_truth, predicted_labels)
    f1 = f1_score(ground_truth, predicted_labels)
    cm = confusion_matrix(ground_truth, predicted_labels)

    doc.add_heading('Metrik Kinerja Utama', level=2)
    
    # Tabel untuk metrik
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Akurasi'
    hdr_cells[1].text = 'Presisi'
    hdr_cells[2].text = 'Recall'
    hdr_cells[3].text = 'F1-Score'
    
    row_cells = table.add_row().cells
    row_cells[0].text = f"{accuracy:.2%}"
    row_cells[1].text = f"{precision:.2%}"
    row_cells[2].text = f"{recall:.2%}"
    row_cells[3].text = f"{f1:.2%}"

    doc.add_heading('Confusion Matrix dan Interpretasi', level=2)
    
    # Buat plot confusion matrix dan simpan ke buffer
    fig, ax = plt.subplots(figsize=(5, 4))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', ax=ax,
                xticklabels=['Prediksi Asli', 'Prediksi Manipulasi'],
                yticklabels=['Aktual Asli', 'Aktual Manipulasi'])
    plt.ylabel('Label Aktual (Ground Truth)')
    plt.xlabel('Label Prediksi Sistem')
    plt.title('Confusion Matrix Kinerja Sistem')
    
    # Simpan plot ke buffer IO
    buf = io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150)
    plt.close(fig)
    buf.seek(0)
    
    # Tambahkan gambar dan interpretasi
    doc.add_picture(buf, width=Inches(4.0))
    doc.add_paragraph(
        f"Interpretasi:\n"
        f"• True Positive (TP): {cm[1, 1]} (Manipulasi terdeteksi benar)\n"
        f"• True Negative (TN): {cm[0, 0]} (Asli terdeteksi benar)\n"
        f"• False Positive (FP): {cm[0, 1]} (Asli dianggap manipulasi)\n"
        f"• False Negative (FN): {cm[1, 0]} (Manipulasi gagal terdeteksi)",
        style='List Bullet'
    )
# ======================= AKHIR FUNGSI BARU =======================

def add_appendix_advanced(doc, analysis_results):
    doc.add_heading('Lampiran A: Rincian Metadata', level=1)
    metadata = analysis_results['metadata']
    meta_str = []
    # Membuat tabel untuk metadata agar lebih rapi
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.style = 'Table Grid'
    meta_table.cell(0, 0).text = 'Tag'
    meta_table.cell(0, 1).text = 'Value'
    
    for key, value in metadata.items():
        if key not in ['Metadata_Inconsistency', 'Metadata_Authenticity_Score']:
            row_cells = meta_table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)

    doc.add_paragraph(f"\nInkonsistensi Metadata Ditemukan: {metadata.get('Metadata_Inconsistency', [])}")
    doc.add_paragraph(f"Skor Keaslian Metadata: {metadata.get('Metadata_Authenticity_Score', 'N/A')}/100")


# ======================= PDF Export Functions =======================

def export_report_pdf(docx_filename, pdf_filename=None):
    """Convert DOCX report to PDF using multiple fallback methods."""
    if not os.path.exists(docx_filename):
        print(f"❌ DOCX file not found: {docx_filename}")
        return None
        
    if pdf_filename is None:
        pdf_filename = docx_filename.replace('.docx', '.pdf')
    
    print(f"📄 Converting DOCX to PDF: {docx_filename} -> {pdf_filename}")
    
    # Method 1: Try using docx2pdf library
    try:
        from docx2pdf import convert
        convert(docx_filename, pdf_filename)
        print(f"📄 PDF report saved as '{pdf_filename}' (via docx2pdf)")
        return pdf_filename
    except (ImportError, Exception) as e:
        print(f"  - docx2pdf failed: {e}. Trying alternative methods...")

    # Method 2: Try using LibreOffice (cross-platform)
    if shutil.which('libreoffice') or shutil.which('soffice'):
        cmd_base = 'libreoffice' if shutil.which('libreoffice') else 'soffice'
        try:
            cmd = [cmd_base, '--headless', '--convert-to', 'pdf', '--outdir',
                   os.path.dirname(os.path.abspath(pdf_filename)) or '.', os.path.abspath(docx_filename)]
            subprocess.run(cmd, check=True, capture_output=True, timeout=60)
            
            generated_pdf_basename = os.path.basename(docx_filename).replace('.docx', '.pdf')
            generated_pdf = os.path.join(os.path.dirname(os.path.abspath(pdf_filename)), generated_pdf_basename)
            
            if os.path.exists(generated_pdf):
                 if os.path.abspath(generated_pdf) != os.path.abspath(pdf_filename):
                    shutil.move(generated_pdf, os.path.abspath(pdf_filename))
                 print(f"📄 PDF report saved as '{pdf_filename}' (via LibreOffice)")
                 return pdf_filename
            else:
                raise FileNotFoundError("LibreOffice did not create the PDF file as expected.")

        except Exception as e:
             print(f"  - LibreOffice failed: {e}. Trying alternative methods...")
    
    # Method 3: Windows-specific (Microsoft Word)
    if platform.system() == 'Windows':
        try:
            import win32com.client as win32
            word = win32.Dispatch('Word.Application')
            word.Visible = False
            doc_path = os.path.abspath(docx_filename)
            pdf_path = os.path.abspath(pdf_filename)
            doc = word.Documents.Open(doc_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
            doc.Close()
            word.Quit()
            print(f"📄 PDF report saved as '{pdf_filename}' (via MS Word)")
            return pdf_filename
        except (ImportError, Exception) as e:
            print(f"  - MS Word COM automation failed: {e}. No more PDF conversion methods available.")

    print("❌ Could not convert DOCX to PDF. Please install one of:")
    print("  - `pip install docx2pdf`")
    print("  - LibreOffice (and ensure it's in your system's PATH)")
    print("  - Microsoft Word (on Windows with `pip install pywin32`)")
    return None